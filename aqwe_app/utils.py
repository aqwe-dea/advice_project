from django.core.mail import send_mail
from docx import Document

def send_advice_email(user_email: str, advice: Advice):
     subject = 'Ваш детальный совет от АКВИ'
     message = f'категория: {advice.category}\n\nВопрос: {advice.question}\n\nОтвет: {advice.answer}\n\nЗаметки: {advice.notes}'
     with open('advice.txt', 'w') as f:
      f.write(message)
     doc = Document()
     doc.add_heading('Ваш детальный совет от АКВИ', level=1)
     doc.add_paragraph(f'Категория: {advice.category}')
     doc.add_paragraph(f'Вопрос: {advice.question}')
     doc.add_paragraph(f'Ответ: {advice.answer}')
     doc.add_paragraph(f'Заметки: {advice.notes}')
     doc.save('advice.doc')
    email = EmailMessage(
         subject = subject,
         body = message,
         from_email = DEFAULT_FROM_EMAIL,
         to = [user_email],
    )
    email.attach_file('advice.txt')
    email.attach_file('advice.doc')
    email.send(fail_silently=False)
       
     subject = 'Ваш совет от АКВИ'
     message = f'Вот ваш совет:\n\n{advice}'
     send_mail(subject, message, DEFAULT_FROM_EMAIL, [user_email], fail_silently=False)